"""Read-only, bounded document ingestion for approved Helper1 folder handles."""
from __future__ import annotations

import hashlib
import io
import multiprocessing
import os
import re
import resource
import socket
import stat
import sys
import time
import unicodedata
import zipfile
from dataclasses import dataclass
from typing import Any, Iterator

from .contracts import canonical_json, hmac_sha256, require_uuid, sha256_bytes
from .parser_isolation import DocumentParserPort, NativeXPCParser
from .security import Helper1SecurityError, open_regular_at, read_bounded_fd

SUPPORTED_EXTENSIONS = frozenset({".docx", ".pdf", ".xlsx", ".md", ".txt"})
MAX_DOCUMENTS = 10_000
MAX_DOCUMENT_BYTES = 64 * 1024 * 1024
MAX_TOTAL_BYTES = 2 * 1024 * 1024 * 1024
MAX_OUTPUT_CHARS = 8_000_000
MAX_ARCHIVE_MEMBERS = 20_000
MAX_ARCHIVE_EXPANDED_BYTES = 256 * 1024 * 1024
MAX_ARCHIVE_RATIO = 200
MAX_XML_BYTES = 32 * 1024 * 1024
DEFAULT_WALL_SECONDS = 20
DEFAULT_RSS_BYTES = 1024 * 1024 * 1024

_ACTIVE_PDF_MARKERS = (
    b"/JavaScript",
    b"/JS",
    b"/Launch",
    b"/EmbeddedFile",
    b"/RichMedia",
    b"/SubmitForm",
    b"/GoToR",
)
_BLOCKED_OOXML_NAMES = (
    "vbaProject.bin",
    "embeddings/",
    "activeX/",
    "customUI/",
)
_HEADING_RE = re.compile(
    r"^(?:#{1,6}\s+.+|(?:제?\s*)?\d+(?:[.-]\d+)*[.)]?\s+\S.+|[^\n]{1,120}\n[-=]{3,})$",
    re.MULTILINE,
)


class IngestionError(RuntimeError):
    pass


@dataclass(frozen=True)
class ParsedDocument:
    source_id: str
    document_revision_sha256: str
    text: str
    media_type: str


@dataclass(frozen=True)
class Chunk:
    chunk_id: str
    source_id: str
    text: str
    ordinal: int
    content_sha256: str

    def to_dict(self) -> dict[str, Any]:
        return {
            "chunk_id": self.chunk_id,
            "source_id": self.source_id,
            "text": self.text,
            "ordinal": self.ordinal,
            "content_sha256": self.content_sha256,
        }


def _file_extension(name: str) -> str:
    lowered = name.casefold()
    for extension in SUPPORTED_EXTENSIONS:
        if lowered.endswith(extension):
            return extension
    return ""


def _walk_regular_files(root_fd: int) -> Iterator[tuple[str, os.stat_result]]:
    stack: list[tuple[int, tuple[str, ...]]] = [(os.dup(root_fd), ())]
    seen_directories: set[tuple[int, int]] = set()
    emitted = 0
    total_bytes = 0
    try:
        while stack:
            directory_fd, prefix = stack.pop()
            try:
                info = os.fstat(directory_fd)
                identity = (info.st_dev, info.st_ino)
                if identity in seen_directories:
                    continue
                seen_directories.add(identity)
                entries = sorted(os.scandir(directory_fd), key=lambda entry: entry.name)
                seen_names: set[str] = set()
                for entry in entries:
                    if entry.name.startswith(".") or entry.name.startswith("~$"):
                        continue
                    canonical_name = unicodedata.normalize("NFKC", entry.name).casefold()
                    if canonical_name in seen_names:
                        raise IngestionError("DOCUMENT_NAME_COLLISION")
                    seen_names.add(canonical_name)
                    if "\x00" in entry.name or entry.name in {".", ".."}:
                        raise IngestionError("DOCUMENT_NAME_INVALID")
                    try:
                        entry_info = entry.stat(follow_symlinks=False)
                    except OSError as exc:
                        raise IngestionError("DOCUMENT_STAT_FAILED") from exc
                    mode = entry_info.st_mode
                    relative_parts = prefix + (entry.name,)
                    relative = "/".join(relative_parts)
                    if stat.S_ISLNK(mode):
                        raise IngestionError("DOCUMENT_SYMLINK_BLOCKED")
                    if stat.S_ISDIR(mode):
                        try:
                            child_fd = os.open(
                                entry.name,
                                os.O_RDONLY
                                | os.O_DIRECTORY
                                | os.O_NOFOLLOW
                                | os.O_CLOEXEC,
                                dir_fd=directory_fd,
                            )
                        except OSError as exc:
                            raise IngestionError("DOCUMENT_DIRECTORY_OPEN_FAILED") from exc
                        stack.append((child_fd, relative_parts))
                        continue
                    if not stat.S_ISREG(mode):
                        raise IngestionError("DOCUMENT_SPECIAL_FILE_BLOCKED")
                    if not _file_extension(relative):
                        continue
                    emitted += 1
                    total_bytes += entry_info.st_size
                    if emitted > MAX_DOCUMENTS:
                        raise IngestionError("DOCUMENT_COUNT_LIMIT")
                    if entry_info.st_size < 0 or entry_info.st_size > MAX_DOCUMENT_BYTES:
                        raise IngestionError("DOCUMENT_SIZE_LIMIT")
                    if total_bytes > MAX_TOTAL_BYTES:
                        raise IngestionError("DOCUMENT_TOTAL_SIZE_LIMIT")
                    yield relative, entry_info
            finally:
                os.close(directory_fd)
    finally:
        for directory_fd, _ in stack:
            os.close(directory_fd)


def _preflight_archive(data: bytes, extension: str) -> None:
    try:
        archive = zipfile.ZipFile(io.BytesIO(data))
    except (zipfile.BadZipFile, OSError) as exc:
        raise IngestionError("DOCUMENT_ARCHIVE_INVALID") from exc
    with archive:
        members = archive.infolist()
        if not members or len(members) > MAX_ARCHIVE_MEMBERS:
            raise IngestionError("DOCUMENT_ARCHIVE_MEMBER_LIMIT")
        expanded = 0
        seen_names: set[str] = set()
        for member in members:
            normalized = member.filename.replace("\\", "/")
            parts = normalized.split("/")
            canonical_name = unicodedata.normalize("NFKC", normalized).casefold()
            if canonical_name in seen_names:
                raise IngestionError("DOCUMENT_ARCHIVE_NAME_COLLISION")
            seen_names.add(canonical_name)
            if (
                normalized.startswith("/")
                or any(part in {"", ".", ".."} for part in parts if part != "")
                or "\x00" in normalized
            ):
                raise IngestionError("DOCUMENT_ARCHIVE_PATH_BLOCKED")
            if member.flag_bits & 0x1:
                raise IngestionError("DOCUMENT_ARCHIVE_ENCRYPTED")
            mode = (member.external_attr >> 16) & 0xFFFF
            if stat.S_ISLNK(mode):
                raise IngestionError("DOCUMENT_ARCHIVE_SYMLINK_BLOCKED")
            expanded += member.file_size
            if expanded > MAX_ARCHIVE_EXPANDED_BYTES:
                raise IngestionError("DOCUMENT_ARCHIVE_EXPANDED_LIMIT")
            compressed = max(1, member.compress_size)
            if member.file_size > 1024 * 1024 and member.file_size / compressed > MAX_ARCHIVE_RATIO:
                raise IngestionError("DOCUMENT_ARCHIVE_RATIO_LIMIT")
            lowered = normalized.casefold()
            if any(blocked.casefold() in lowered for blocked in _BLOCKED_OOXML_NAMES):
                raise IngestionError("DOCUMENT_ACTIVE_CONTENT_BLOCKED")
            if extension == ".xlsx" and lowered.endswith((".xlsm", ".xlam", ".bin")):
                raise IngestionError("DOCUMENT_MACRO_BLOCKED")
            if lowered.endswith((".xml", ".rels")):
                if member.file_size > MAX_XML_BYTES:
                    raise IngestionError("DOCUMENT_XML_SIZE_LIMIT")
                xml = archive.read(member)
                upper = xml.upper()
                if b"<!DOCTYPE" in upper or b"<!ENTITY" in upper:
                    raise IngestionError("DOCUMENT_XML_ENTITY_BLOCKED")
                if lowered.endswith(".rels") and (
                    b'TARGETMODE="EXTERNAL"' in upper
                    or b"TARGETMODE='EXTERNAL'" in upper
                ):
                    raise IngestionError("DOCUMENT_EXTERNAL_RELATIONSHIP_BLOCKED")


def _parse_bytes(data: bytes, extension: str) -> str:
    if extension == ".pdf" and not data.startswith(b"%PDF-"):
        raise IngestionError("DOCUMENT_MAGIC_MISMATCH")
    if extension in {".docx", ".xlsx"} and not data.startswith(b"PK\x03\x04"):
        raise IngestionError("DOCUMENT_MAGIC_MISMATCH")
    if extension in {".txt", ".md"} and b"\x00" in data:
        raise IngestionError("DOCUMENT_MAGIC_MISMATCH")
    if extension in {".txt", ".md"}:
        try:
            return data.decode("utf-8-sig")
        except UnicodeDecodeError as exc:
            raise IngestionError("DOCUMENT_ENCODING_UNSUPPORTED") from exc
    if extension == ".pdf":
        if any(marker in data for marker in _ACTIVE_PDF_MARKERS):
            raise IngestionError("DOCUMENT_PDF_ACTIVE_CONTENT_BLOCKED")
        try:
            from pdfminer.high_level import extract_text

            return extract_text(io.BytesIO(data))
        except Exception as exc:
            raise IngestionError("DOCUMENT_PDF_PARSE_FAILED") from exc
    if extension == ".docx":
        _preflight_archive(data, extension)
        try:
            from docx import Document

            document = Document(io.BytesIO(data))
            blocks: list[str] = []
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    blocks.append(paragraph.text)
            for table in document.tables:
                for row in table.rows:
                    values = [cell.text.strip() for cell in row.cells]
                    if any(values):
                        blocks.append("\t".join(values))
            return "\n".join(blocks)
        except IngestionError:
            raise
        except Exception as exc:
            raise IngestionError("DOCUMENT_DOCX_PARSE_FAILED") from exc
    if extension == ".xlsx":
        _preflight_archive(data, extension)
        try:
            import openpyxl

            workbook = openpyxl.load_workbook(
                io.BytesIO(data),
                read_only=True,
                data_only=True,
                keep_links=False,
            )
            blocks = []
            for sheet in workbook.worksheets:
                if sheet.sheet_state != "visible":
                    continue
                blocks.append(f"# {sheet.title}")
                for row_index, row in enumerate(sheet.iter_rows(), start=1):
                    if sheet.row_dimensions[row_index].hidden:
                        continue
                    values = [
                        "" if cell.value is None else str(cell.value)
                        for cell in row
                        if not sheet.column_dimensions[cell.column_letter].hidden
                    ]
                    if any(values):
                        blocks.append("\t".join(values))
            workbook.close()
            return "\n".join(blocks)
        except IngestionError:
            raise
        except Exception as exc:
            raise IngestionError("DOCUMENT_XLSX_PARSE_FAILED") from exc
    raise IngestionError("DOCUMENT_TYPE_UNSUPPORTED")


def _deny_network(*_args: Any, **_kwargs: Any) -> Any:
    raise OSError("NETWORK_DISABLED")


def _parser_child(
    connection: multiprocessing.connection.Connection,
    data: bytes,
    extension: str,
    rss_bytes: int,
) -> None:
    try:
        resource.setrlimit(resource.RLIMIT_CPU, (10, 10))
        try:
            if hasattr(resource, "RLIMIT_RSS"):
                resource.setrlimit(resource.RLIMIT_RSS, (rss_bytes, rss_bytes))
        except (OSError, ValueError):
            pass
        resource.setrlimit(resource.RLIMIT_FSIZE, (0, 0))
        socket.socket = _deny_network  # type: ignore[assignment]
        socket.create_connection = _deny_network  # type: ignore[assignment]
        text = _parse_bytes(data, extension)
        if len(text) > MAX_OUTPUT_CHARS:
            raise IngestionError("DOCUMENT_OUTPUT_LIMIT")
        peak_rss = resource.getrusage(resource.RUSAGE_SELF).ru_maxrss
        if sys.platform != "darwin":
            peak_rss *= 1024
        if peak_rss > rss_bytes:
            raise IngestionError("DOCUMENT_PARSER_RSS_LIMIT")
        connection.send(("ok", text, peak_rss))
    except BaseException as exc:
        code = str(exc)
        if re.fullmatch(r"[A-Z][A-Z0-9_]{2,95}", code) is None:
            code = "DOCUMENT_PARSE_FAILED_" + exc.__class__.__name__.upper()
        try:
            connection.send(("error", code))
        except Exception:
            pass
    finally:
        connection.close()


def parse_in_sandbox(
    data: bytes,
    extension: str,
    *,
    wall_seconds: int = DEFAULT_WALL_SECONDS,
    rss_bytes: int = DEFAULT_RSS_BYTES,
) -> str:
    if extension not in SUPPORTED_EXTENSIONS:
        raise IngestionError("DOCUMENT_TYPE_UNSUPPORTED")
    if not data or len(data) > MAX_DOCUMENT_BYTES:
        raise IngestionError("DOCUMENT_SIZE_LIMIT")
    context = multiprocessing.get_context("spawn")
    parent, child = context.Pipe(duplex=False)
    process = context.Process(
        target=_parser_child,
        args=(child, data, extension, rss_bytes),
        daemon=True,
    )
    process.start()
    child.close()
    deadline = time.monotonic() + wall_seconds
    try:
        while time.monotonic() < deadline:
            if parent.poll(0.05):
                message = parent.recv()
                status, value = message[0], message[1]
                process.join(timeout=1)
                if status != "ok":
                    raise IngestionError(value)
                if not isinstance(value, str) or not value.strip():
                    raise IngestionError("DOCUMENT_TEXT_EMPTY")
                return value
            if not process.is_alive():
                raise IngestionError("DOCUMENT_PARSER_EXITED")
        process.kill()
        process.join(timeout=1)
        raise IngestionError("DOCUMENT_PARSER_TIMEOUT")
    finally:
        parent.close()
        if process.is_alive():
            process.kill()
            process.join(timeout=1)


@dataclass
class ApprovedFolderIngestor:
    workspace_id: str
    folder_fd: int
    workspace_key: bytes
    parser: DocumentParserPort | None = None
    allow_test_parser: bool = False

    def __post_init__(self) -> None:
        require_uuid(self.workspace_id, "WORKSPACE_ID_INVALID")
        if len(self.workspace_key) != 32:
            raise IngestionError("WORKSPACE_KEY_INVALID")
        info = os.fstat(self.folder_fd)
        if not stat.S_ISDIR(info.st_mode):
            raise IngestionError("FOLDER_HANDLE_INVALID")
        if type(self.allow_test_parser) is not bool:
            raise IngestionError("PARSER_MODE_INVALID")
        if self.parser is None and not self.allow_test_parser:
            raise IngestionError("REFUSED_PARSER_UNAVAILABLE")
        if self.parser is not None and (
            type(self.parser) is not NativeXPCParser
            or not self.parser.is_production_parser
        ):
            raise IngestionError("PARSER_PRODUCTION_BOUNDARY_REQUIRED")

    def read_documents(self) -> tuple[ParsedDocument, ...]:
        documents: list[ParsedDocument] = []
        for relative, expected in _walk_regular_files(self.folder_fd):
            fd = open_regular_at(self.folder_fd, relative)
            try:
                actual = os.fstat(fd)
                if (
                    actual.st_dev != expected.st_dev
                    or actual.st_ino != expected.st_ino
                    or actual.st_size != expected.st_size
                ):
                    raise IngestionError("DOCUMENT_CHANGED_DURING_READ")
                data = read_bounded_fd(fd, maximum=MAX_DOCUMENT_BYTES)
            except Helper1SecurityError as exc:
                raise IngestionError(str(exc)) from exc
            finally:
                os.close(fd)
            revision = sha256_bytes(data)
            source_id = "src:" + hmac_sha256(
                self.workspace_key,
                b"butler/helper1/source/v2",
                relative.encode("utf-8"),
                revision.encode("ascii"),
            )[:48]
            extension = _file_extension(relative)
            text = (
                self.parser.parse(data, extension)
                if self.parser is not None
                else parse_in_sandbox(data, extension)
            )
            documents.append(
                ParsedDocument(
                    source_id=source_id,
                    document_revision_sha256=revision,
                    text=text,
                    media_type=extension[1:],
                )
            )
        if not documents:
            raise IngestionError("DOCUMENT_SET_EMPTY")
        return tuple(documents)


@dataclass(frozen=True)
class HeadingChunker:
    max_chars: int = 1800
    overlap_chars: int = 120
    policy_version: str = "heading-aware-char-v2"

    def __post_init__(self) -> None:
        if isinstance(self.max_chars, bool) or not 256 <= self.max_chars <= 8192:
            raise IngestionError("CHUNK_MAX_INVALID")
        if (
            isinstance(self.overlap_chars, bool)
            or self.overlap_chars < 0
            or self.overlap_chars >= self.max_chars // 2
        ):
            raise IngestionError("CHUNK_OVERLAP_INVALID")

    @property
    def policy_sha256(self) -> str:
        return sha256_bytes(canonical_json({
            "policy_version": self.policy_version,
            "unicode_normalization": "NFKC",
            "line_endings": "LF",
            "whitespace_folding": "strip_chunk_edges",
            "heading_rules": "heading-regex-v2",
            "max_chars": self.max_chars,
            "overlap_chars": self.overlap_chars,
            "oversize_block": "sliding-window",
        }))

    def _segments(self, text: str) -> list[str]:
        normalized = unicodedata.normalize("NFKC", text).replace("\r\n", "\n").replace("\r", "\n")
        lines = [line.rstrip() for line in normalized.split("\n")]
        segments: list[str] = []
        current: list[str] = []
        for line in lines:
            candidate = line.strip()
            is_heading = bool(candidate and _HEADING_RE.fullmatch(candidate))
            if is_heading and current:
                value = "\n".join(current).strip()
                if value:
                    segments.append(value)
                current = [candidate]
            else:
                current.append(line)
        value = "\n".join(current).strip()
        if value:
            segments.append(value)
        return segments

    def chunk(
        self,
        documents: tuple[ParsedDocument, ...],
        workspace_id: str,
        workspace_key: bytes,
    ) -> tuple[Chunk, ...]:
        require_uuid(workspace_id, "WORKSPACE_ID_INVALID")
        chunks: list[Chunk] = []
        for document in documents:
            ordinal = 0
            for segment in self._segments(document.text):
                start = 0
                while start < len(segment):
                    end = min(len(segment), start + self.max_chars)
                    value = segment[start:end].strip()
                    if value:
                        content_digest = sha256_bytes(value.encode("utf-8"))
                        chunk_id = "chk:" + hmac_sha256(
                            workspace_key,
                            b"butler/helper1/chunk/v2",
                            workspace_id.encode("ascii"),
                            document.source_id.encode("ascii"),
                            document.document_revision_sha256.encode("ascii"),
                            ordinal.to_bytes(8, "big"),
                            content_digest.encode("ascii"),
                            self.policy_sha256.encode("ascii"),
                        )[:48]
                        chunks.append(
                            Chunk(
                                chunk_id=chunk_id,
                                source_id=document.source_id,
                                text=value,
                                ordinal=ordinal,
                                content_sha256=content_digest,
                            )
                        )
                        ordinal += 1
                    if end == len(segment):
                        break
                    start = end - self.overlap_chars
        if not chunks:
            raise IngestionError("CHUNK_SET_EMPTY")
        return tuple(chunks)
