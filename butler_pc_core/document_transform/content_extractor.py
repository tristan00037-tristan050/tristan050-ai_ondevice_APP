"""content_extractor.py — 외부 문서 텍스트 추출 (.docx / .pdf / .txt / .md / .eml)."""
from __future__ import annotations

import email
import io
from pathlib import Path


def extract_text_from_file(path: Path | str) -> str:
    """파일 경로에서 텍스트 추출."""
    path = Path(path)
    suffix = path.suffix.lower()
    data = path.read_bytes()
    return extract_text_from_bytes(data, suffix)


def extract_text_from_bytes(data: bytes, suffix: str) -> str:
    """바이트 + 확장자에서 텍스트 추출."""
    suffix = suffix.lower().lstrip(".")
    if suffix in ("txt", "md"):
        return data.decode("utf-8", errors="replace")
    if suffix == "docx":
        return _extract_docx(data)
    if suffix == "pdf":
        return _extract_pdf(data)
    if suffix == "eml":
        return _extract_eml(data)
    raise ValueError(f"지원하지 않는 파일 형식: .{suffix} (지원: .txt .md .docx .pdf .eml)")


def _extract_docx(data: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        raise RuntimeError("python-docx 미설치 (pip install python-docx)")
    doc = docx.Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _extract_pdf(data: bytes) -> str:
    try:
        from pdfminer.high_level import extract_text_to_fp
        from pdfminer.layout import LAParams
    except ImportError:
        raise RuntimeError("pdfminer.six 미설치 (pip install pdfminer.six)")
    out = io.StringIO()
    extract_text_to_fp(io.BytesIO(data), out, laparams=LAParams())
    return out.getvalue()


def _extract_eml(data: bytes) -> str:
    msg = email.message_from_bytes(data)
    parts: list[str] = []
    subject = msg.get("Subject", "")
    if subject:
        parts.append(f"제목: {subject}")
    sender = msg.get("From", "")
    if sender:
        parts.append(f"발신: {sender}")
    if msg.is_multipart():
        for part in msg.walk():
            if part.get_content_type() == "text/plain":
                payload = part.get_payload(decode=True)
                if payload:
                    parts.append(payload.decode("utf-8", errors="replace"))
    else:
        payload = msg.get_payload(decode=True)
        if payload:
            parts.append(payload.decode("utf-8", errors="replace"))
    return "\n".join(parts)
