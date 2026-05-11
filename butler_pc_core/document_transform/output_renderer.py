"""output_renderer.py — 매핑 결과를 .docx / .md로 렌더링."""
from __future__ import annotations

import io
from typing import List


def render_docx_bytes(sections: List[dict]) -> bytes:
    """
    sections: [{"heading": str, "level": int, "content": str, "mapped": bool}, ...]
    → .docx bytes
    """
    try:
        import docx
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx 미설치 (pip install python-docx)")

    doc = docx.Document()

    heading_style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}

    for sec in sections:
        heading = sec.get("heading", "")
        level = sec.get("level", 1)
        content = sec.get("content", "")
        mapped = sec.get("mapped", True)

        style_name = heading_style_map.get(level, "Heading 2")
        doc.add_paragraph(heading, style=style_name)

        if content:
            para = doc.add_paragraph(content)
        else:
            para = doc.add_paragraph()
            run = para.add_run("(원본에 해당 내용 없음)")
            run.italic = True
            if not mapped:
                run.font.color.rgb = None  # keep default

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_md_text(sections: List[dict]) -> str:
    """sections → .md text."""
    lines: list[str] = []
    for sec in sections:
        heading = sec.get("heading", "")
        level = sec.get("level", 1)
        content = sec.get("content", "")
        mapped = sec.get("mapped", True)

        prefix = "#" * level
        lines.append(f"{prefix} {heading}")
        lines.append("")
        if content:
            lines.append(content)
        else:
            note = "_(원본에 해당 내용 없음)_" if not mapped else "_(내용 없음)_"
            lines.append(note)
        lines.append("")

    return "\n".join(lines)
