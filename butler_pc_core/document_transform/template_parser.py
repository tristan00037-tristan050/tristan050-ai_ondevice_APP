"""template_parser.py — 우리 양식 구조 분석 (.docx / .md)."""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List


@dataclass
class TemplateSection:
    heading: str
    level: int           # 1=h1, 2=h2, 3=h3
    placeholder_text: str
    format_hint: str     # "paragraph" | "table" | "bullet_list"
    is_required: bool = True


def parse_template_file(path: Path | str) -> List[TemplateSection]:
    """파일 경로에서 양식 구조 추출."""
    path = Path(path)
    suffix = path.suffix.lower()
    data = path.read_bytes()
    return parse_template_bytes(data, suffix)


def parse_template_bytes(data: bytes, suffix: str) -> List[TemplateSection]:
    """바이트 + 확장자에서 양식 구조 추출."""
    suffix = suffix.lower().lstrip(".")
    if suffix == "docx":
        return _parse_docx_template(data)
    if suffix == "md":
        return _parse_md_template(data.decode("utf-8", errors="replace"))
    raise ValueError(f"양식 파일은 .docx 또는 .md만 지원합니다 (.{suffix} 불가)")


def _parse_docx_template(data: bytes) -> List[TemplateSection]:
    try:
        import docx
    except ImportError:
        raise RuntimeError("python-docx 미설치 (pip install python-docx)")

    doc = docx.Document(io.BytesIO(data))
    sections: List[TemplateSection] = []
    current_heading: str | None = None
    current_level: int = 1
    body_parts: list[str] = []
    has_table = False
    has_list = False

    def _flush():
        nonlocal current_heading, body_parts, has_table, has_list
        if current_heading is None:
            return
        if has_table:
            fmt = "table"
        elif has_list:
            fmt = "bullet_list"
        else:
            fmt = "paragraph"
        sections.append(TemplateSection(
            heading=current_heading,
            level=current_level,
            placeholder_text="\n".join(body_parts),
            format_hint=fmt,
            is_required=True,
        ))
        body_parts = []
        has_table = False
        has_list = False

    heading_styles = {"heading 1": 1, "heading 2": 2, "heading 3": 3}

    for para in doc.paragraphs:
        style_name = para.style.name.lower() if para.style else ""
        level = heading_styles.get(style_name)
        if level is not None and para.text.strip():
            _flush()
            current_heading = para.text.strip()
            current_level = level
        elif current_heading is not None and para.text.strip():
            text = para.text.strip()
            is_list = style_name.startswith("list") or text.startswith(("•", "-", "*", "·"))
            if is_list:
                has_list = True
            body_parts.append(text)

    for table in doc.tables:
        has_table = True
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                body_parts.append(row_text)

    _flush()

    if not sections:
        sections.append(TemplateSection(
            heading="본문",
            level=1,
            placeholder_text="",
            format_hint="paragraph",
            is_required=True,
        ))
    return sections


def _parse_md_template(text: str) -> List[TemplateSection]:
    sections: List[TemplateSection] = []
    heading_re = re.compile(r"^(#{1,3})\s+(.+)$")
    list_re = re.compile(r"^\s*[-*•]\s+")

    current_heading: str | None = None
    current_level: int = 1
    body_lines: list[str] = []

    def _flush():
        nonlocal current_heading, body_lines
        if current_heading is None:
            return
        body = "\n".join(body_lines).strip()
        has_list = any(list_re.match(l) for l in body_lines)
        fmt = "bullet_list" if has_list else "paragraph"
        sections.append(TemplateSection(
            heading=current_heading,
            level=current_level,
            placeholder_text=body,
            format_hint=fmt,
            is_required=True,
        ))
        body_lines.clear()

    for line in text.splitlines():
        m = heading_re.match(line)
        if m:
            _flush()
            current_level = len(m.group(1))
            current_heading = m.group(2).strip()
        elif current_heading is not None:
            body_lines.append(line)

    _flush()

    if not sections:
        sections.append(TemplateSection(
            heading="본문",
            level=1,
            placeholder_text=text.strip(),
            format_hint="paragraph",
            is_required=True,
        ))
    return sections
