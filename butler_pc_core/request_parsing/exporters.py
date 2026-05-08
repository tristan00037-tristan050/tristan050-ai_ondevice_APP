"""exporters.py — ParsedResult → Markdown / docx bytes."""
from __future__ import annotations

from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from .parser import ParsedResult

_PRIORITY_LABEL = {"P1": "🔴 P1 (긴급·필수)", "P2": "🟠 P2 (권장)", "P3": "⚪ P3 (선택)"}


def result_to_markdown(result: "ParsedResult", original_text: str = "") -> str:
    """ParsedResult → Markdown 문자열."""
    from .parser import confidence_level

    lines: list[str] = [
        "# 요청 파싱 결과",
        "",
        f"**신뢰도**: {result.confidence:.0%} ({confidence_level(result.confidence).upper()})",
        f"**입력 형식**: {result.input_format}",
        "",
    ]

    # 의도
    lines += [
        "## 발신자 의도",
        result.intent.summary or "(분석 중)",
        f"- 톤: {result.intent.tone}",
        f"- 기대 응답: {result.intent.expected_response or '—'}",
        "",
    ]

    # 마감일
    lines += ["## 마감일"]
    if result.deadline.raw_text:
        lines.append(f"- 원문: {result.deadline.raw_text}")
        lines.append(
            f"- 파싱 날짜: {result.deadline.parsed_date or '날짜 특정 불가'}"
        )
    else:
        lines.append("- 마감일 언급 없음")
    lines.append("")

    # 액션 목록
    lines += ["## 액션 목록"]
    if result.actions:
        for i, action in enumerate(result.actions, 1):
            label = _PRIORITY_LABEL.get(action.priority, action.priority)
            lines.append(f"{i}. [{label}] {action.text}")
            if action.rationale:
                lines.append(f"   - 근거: {action.rationale}")
    else:
        lines.append("- 명시적 액션 없음")
    lines.append("")

    # 필요 자료
    lines += ["## 필요 자료"]
    if result.required_materials:
        for mat in result.required_materials:
            optional_tag = " *(선택)*" if mat.is_optional else ""
            lines.append(f"- {mat.name}{optional_tag}")
            if mat.rationale:
                lines.append(f"  - {mat.rationale}")
    else:
        lines.append("- 필요 자료 없음")
    lines.append("")

    # 원문 (PII 마스킹 후)
    if result.masked_text:
        lines += [
            "## 원문 (PII 마스킹)",
            "```",
            result.masked_text[:2000],
            "```",
            "",
        ]

    return "\n".join(lines)


def result_to_docx_bytes(result: "ParsedResult", original_text: str = "") -> bytes:
    """ParsedResult → .docx 바이트. python-docx 없으면 ParseError."""
    try:
        import docx
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        from .parser import ParseError
        raise ParseError("python-docx 미설치 — pip install python-docx")

    import io
    from .parser import confidence_level

    doc = docx.Document()

    # 제목
    title = doc.add_heading("요청 파싱 결과", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 메타
    meta = doc.add_paragraph()
    meta.add_run("신뢰도: ").bold = True
    level = confidence_level(result.confidence)
    meta.add_run(f"{result.confidence:.0%} ({level.upper()})")

    # 의도
    doc.add_heading("발신자 의도", level=2)
    doc.add_paragraph(result.intent.summary or "(분석 중)")
    p = doc.add_paragraph()
    p.add_run("톤: ").bold = True
    p.add_run(result.intent.tone)
    if result.intent.expected_response:
        p2 = doc.add_paragraph()
        p2.add_run("기대 응답: ").bold = True
        p2.add_run(result.intent.expected_response)

    # 마감일
    doc.add_heading("마감일", level=2)
    if result.deadline.raw_text:
        p = doc.add_paragraph()
        p.add_run("원문: ").bold = True
        p.add_run(result.deadline.raw_text)
        p2 = doc.add_paragraph()
        p2.add_run("파싱 날짜: ").bold = True
        p2.add_run(result.deadline.parsed_date or "날짜 특정 불가")
    else:
        doc.add_paragraph("마감일 언급 없음")

    # 액션 목록
    doc.add_heading("액션 목록", level=2)
    if result.actions:
        for i, action in enumerate(result.actions, 1):
            label = _PRIORITY_LABEL.get(action.priority, action.priority)
            p = doc.add_paragraph(style="List Number")
            p.add_run(f"[{label}] {action.text}")
            if action.rationale:
                doc.add_paragraph(f"근거: {action.rationale}", style="List Bullet")
    else:
        doc.add_paragraph("명시적 액션 없음")

    # 필요 자료
    doc.add_heading("필요 자료", level=2)
    if result.required_materials:
        for mat in result.required_materials:
            optional_tag = " (선택)" if mat.is_optional else ""
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{mat.name}{optional_tag}")
            if mat.rationale:
                doc.add_paragraph(mat.rationale, style="List Bullet 2")
    else:
        doc.add_paragraph("필요 자료 없음")

    # 원문 (마스킹)
    if result.masked_text:
        doc.add_heading("원문 (PII 마스킹)", level=2)
        doc.add_paragraph(result.masked_text[:2000])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
