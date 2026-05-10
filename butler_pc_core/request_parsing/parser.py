"""parser.py — D-3 카드 1 요청 핵심 파악·정리 핵심 로직

PII 마스킹, 한국어 날짜 추출, JSON Schema 검증, LLM 기반 파싱
"""
from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from datetime import date, timedelta
from pathlib import Path
from typing import Any, Optional

# ── 입력 길이 제약 §6.3 ───────────────────────────────────────────────────────
MIN_TEXT_LENGTH = 30
MAX_TEXT_LENGTH = 32_000

# ── PII 패턴 (D-2 anonymize.py 패턴 동일) ────────────────────────────────────
_PII_PATTERNS: dict[str, re.Pattern[str]] = {
    "rrn":     re.compile(r"\d{6}-[1-4]\d{6}"),
    "biz_no":  re.compile(r"\d{3}-\d{2}-\d{5}"),
    "card":    re.compile(r"\d{4}[- ]\d{4}[- ]\d{4}[- ]\d{4}"),
    "account": re.compile(r"\d{2,6}-\d{2,6}-\d{2,6}(?:-\d{2,6})?"),
    "phone":   re.compile(r"01[016789]-?\d{3,4}-?\d{4}"),
    "phone_local": re.compile(r"0[2-9]\d-?\d{3,4}-?\d{4}"),
    "email":   re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}"),
}


# ── 예외 계층 ────────────────────────────────────────────────────────────────

class ParseError(Exception):
    """요청 파싱 기본 오류."""


class TextTooShortError(ParseError):
    """텍스트가 너무 짧음 (30자 미만)."""


class TextTooLongError(ParseError):
    """텍스트가 너무 김 (32,000자 초과)."""


# ── 데이터 모델 ───────────────────────────────────────────────────────────────

@dataclass
class ActionItem:
    text: str
    priority: str = "P2"    # P1 | P2 | P3
    rationale: str = ""


@dataclass
class Deadline:
    raw_text: str = ""
    parsed_date: Optional[str] = None  # ISO 8601 "YYYY-MM-DD" or None
    confidence: float = 0.5
    time_text: str = ""  # e.g. "오후 3시", "오전 10시 30분"


@dataclass
class MaterialItem:
    name: str
    is_optional: bool = False
    rationale: str = ""


@dataclass
class Intent:
    summary: str = ""
    tone: str = "formal"   # formal | informal | urgent | casual
    expected_response: str = ""


@dataclass
class ParsedResult:
    actions: list[ActionItem] = field(default_factory=list)
    deadline: Deadline = field(default_factory=Deadline)
    required_materials: list[MaterialItem] = field(default_factory=list)
    intent: Intent = field(default_factory=Intent)
    confidence: float = 0.75
    masked_text: str = ""
    input_format: str = "text"

    def to_dict(self) -> dict:
        return {
            "actions": [
                {"text": a.text, "priority": a.priority, "rationale": a.rationale}
                for a in self.actions
            ],
            "deadline": {
                "raw_text": self.deadline.raw_text,
                "parsed_date": self.deadline.parsed_date,
                "confidence": self.deadline.confidence,
                "time_text": self.deadline.time_text,
            },
            "required_materials": [
                {"name": m.name, "is_optional": m.is_optional, "rationale": m.rationale}
                for m in self.required_materials
            ],
            "intent": {
                "summary": self.intent.summary,
                "tone": self.intent.tone,
                "expected_response": self.intent.expected_response,
            },
            "confidence": self.confidence,
        }

    @classmethod
    def from_dict(cls, d: dict) -> "ParsedResult":
        result = cls()
        result.actions = [
            ActionItem(
                text=a.get("text", ""),
                priority=a.get("priority", "P2"),
                rationale=a.get("rationale", ""),
            )
            for a in d.get("actions", [])
        ]
        dl = d.get("deadline") or {}
        result.deadline = Deadline(
            raw_text=dl.get("raw_text", ""),
            parsed_date=dl.get("parsed_date"),
            confidence=float(dl.get("confidence", 0.5)),
            time_text=dl.get("time_text", ""),
        )
        result.required_materials = [
            MaterialItem(
                name=m.get("name", ""),
                is_optional=bool(m.get("is_optional", False)),
                rationale=m.get("rationale", ""),
            )
            for m in d.get("required_materials", [])
        ]
        intent = d.get("intent", {})
        result.intent = Intent(
            summary=intent.get("summary", ""),
            tone=intent.get("tone", "formal"),
            expected_response=intent.get("expected_response", ""),
        )
        result.confidence = float(d.get("confidence", 0.75))
        result.masked_text = d.get("masked_text", "")
        result.input_format = d.get("input_format", "text")
        return result


# ── PII 마스킹 ────────────────────────────────────────────────────────────────

def mask_pii(text: str) -> str:
    """개인정보 패턴을 마스킹 태그로 치환."""
    for pii_type, pattern in _PII_PATTERNS.items():
        text = pattern.sub(f"<{pii_type}_masked>", text)
    return text


# ── 한국어 날짜 추출 §4.4 ─────────────────────────────────────────────────────

def parse_korean_date(text: str, today: Optional[date] = None) -> Optional[str]:
    """
    한국어 날짜 표현 → ISO 8601 문자열 ("YYYY-MM-DD") 또는 None.

    지원 패턴:
    - "다음 주 화요일" → 다음 주 화요일
    - "이번 주 금요일" → 이번 주 금요일
    - "다다음 주" → 다다음 주 월요일
    - "이번 달 말" / "월말" → 이번 달 마지막 날
    - "내일" → 내일
    - "모레" → 모레
    - "N월 M일" → 해당 날짜
    - "N일" / "M일까지" → 이번/다음 달 해당 일
    """
    if today is None:
        today = date.today()

    KO_WEEKDAY = {"월": 0, "화": 1, "수": 2, "목": 3, "금": 4, "토": 5, "일": 6}

    # 내일 / 모레
    if re.search(r"내일", text):
        return (today + timedelta(days=1)).isoformat()
    if re.search(r"모레", text):
        return (today + timedelta(days=2)).isoformat()

    # "다다음 주" → +2주 월요일
    if re.search(r"다다음\s*주", text):
        days_to_monday = (7 - today.weekday()) % 7 or 7
        next_monday = today + timedelta(days=days_to_monday)
        return (next_monday + timedelta(weeks=1)).isoformat()

    # "다음 주 요일" 또는 "다음주 요일"
    m = re.search(r"다음\s*주\s*([월화수목금토일])요일?", text)
    if m:
        target_wd = KO_WEEKDAY[m.group(1)]
        # 다음 주 시작(월요일) 기준
        days_to_next_monday = (7 - today.weekday()) % 7 or 7
        next_monday = today + timedelta(days=days_to_next_monday)
        delta = (target_wd - next_monday.weekday()) % 7
        return (next_monday + timedelta(days=delta)).isoformat()

    # "다음 주" (요일 없음) → 다음 주 월요일
    if re.search(r"다음\s*주", text):
        days_to_next_monday = (7 - today.weekday()) % 7 or 7
        return (today + timedelta(days=days_to_next_monday)).isoformat()

    # "이번 주 요일"
    m = re.search(r"이번\s*주\s*([월화수목금토일])요일?", text)
    if m:
        target_wd = KO_WEEKDAY[m.group(1)]
        delta = (target_wd - today.weekday()) % 7
        if delta == 0 and today.weekday() != target_wd:
            delta = 7
        return (today + timedelta(days=delta)).isoformat()

    # "이번 달 말" / "월말" / "이번 달 마지막"
    if re.search(r"이번\s*달\s*(말|마지막)|월말", text):
        import calendar
        last_day = calendar.monthrange(today.year, today.month)[1]
        return date(today.year, today.month, last_day).isoformat()

    # "다음 달 말"
    if re.search(r"다음\s*달\s*(말|마지막)", text):
        import calendar
        next_month = today.month % 12 + 1
        year = today.year + (1 if today.month == 12 else 0)
        last_day = calendar.monthrange(year, next_month)[1]
        return date(year, next_month, last_day).isoformat()

    # "N월 M일" — 절대 날짜
    m = re.search(r"(\d{1,2})월\s*(\d{1,2})일", text)
    if m:
        month, day = int(m.group(1)), int(m.group(2))
        try:
            candidate = date(today.year, month, day)
            if candidate < today:
                candidate = date(today.year + 1, month, day)
            return candidate.isoformat()
        except ValueError:
            pass

    # "M일" / "M일까지" — 이번 달 or 다음 달
    m = re.search(r"(?<!\d)(\d{1,2})일(?:까지)?(?!\d)", text)
    if m:
        day = int(m.group(1))
        try:
            candidate = date(today.year, today.month, day)
            if candidate < today:
                next_month = today.month % 12 + 1
                year = today.year + (1 if today.month == 12 else 0)
                candidate = date(year, next_month, day)
            return candidate.isoformat()
        except ValueError:
            pass

    # ISO 날짜 직접 언급 "YYYY-MM-DD" 또는 "YYYY.MM.DD" 또는 "YYYY년 MM월 DD일"
    m = re.search(r"(\d{4})[-./년]\s*(\d{1,2})[-./월]\s*(\d{1,2})일?", text)
    if m:
        try:
            return date(int(m.group(1)), int(m.group(2)), int(m.group(3))).isoformat()
        except ValueError:
            pass

    return None


# ── 한국어 시간 추출 ──────────────────────────────────────────────────────────

def parse_korean_time(text: str) -> str:
    """한국어 시간 표현 → 문자열. 예: "오후 3시", "오전 10시 30분"."""
    m = re.search(r"(오전|오후|낮|밤|새벽)?\s*(\d{1,2})시(?:\s*(\d{1,2})분)?", text)
    if not m:
        return ""
    period = m.group(1) or ""
    hour = m.group(2)
    minute = m.group(3)
    result = f"{period} {hour}시".strip() if period else f"{hour}시"
    if minute:
        result += f" {minute}분"
    return result


# ── JSON Schema 검증 ──────────────────────────────────────────────────────────

_REQUIRED_KEYS = {"actions", "deadline", "required_materials", "intent"}
_PRIORITY_VALUES = {"P1", "P2", "P3"}
_TONE_VALUES = {"formal", "informal", "urgent", "casual"}


def validate_schema(data: dict) -> list[str]:
    """ParsedResult dict에 대한 스키마 검증. 오류 메시지 리스트 반환 (비면 통과)."""
    errors: list[str] = []

    missing = _REQUIRED_KEYS - set(data.keys())
    if missing:
        errors.append(f"필수 키 누락: {missing}")

    for action in data.get("actions", []):
        if "text" not in action:
            errors.append("actions.text 필수")
        p = action.get("priority", "P2")
        if p not in _PRIORITY_VALUES:
            errors.append(f"actions.priority 유효하지 않음: {p}")

    dl = data.get("deadline", {})
    if "raw_text" not in dl or "parsed_date" not in dl:
        errors.append("deadline.raw_text + parsed_date 필수")

    for mat in data.get("required_materials", []):
        if "name" not in mat:
            errors.append("required_materials.name 필수")

    intent = data.get("intent", {})
    if "summary" not in intent or "tone" not in intent:
        errors.append("intent.summary + tone 필수")
    tone = intent.get("tone", "formal")
    if tone not in _TONE_VALUES:
        errors.append(f"intent.tone 유효하지 않음: {tone}")

    conf = data.get("confidence")
    if conf is not None and not (0.0 <= float(conf) <= 1.0):
        errors.append(f"confidence 범위 초과: {conf}")

    return errors


# ── 입력 형식별 텍스트 추출 §3.3 ─────────────────────────────────────────────

def extract_text_from_file(file_path: str, suffix: str) -> str:
    """파일 → 텍스트 추출. suffix는 .txt / .docx / .pdf / .md / .eml."""
    p = Path(file_path)
    suffix = suffix.lower()

    if suffix in (".txt", ".md"):
        for enc in ("utf-8-sig", "utf-8", "cp949", "euc-kr"):
            try:
                return p.read_text(encoding=enc)
            except UnicodeDecodeError:
                continue
        return p.read_bytes().decode("utf-8", errors="replace")

    if suffix == ".docx":
        try:
            import docx  # python-docx
            doc = docx.Document(str(p))
            return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        except ImportError:
            raise ParseError("python-docx 미설치 — pip install python-docx")

    if suffix == ".pdf":
        try:
            from pdfminer.high_level import extract_text as _pdf_extract
            return _pdf_extract(str(p)) or ""
        except ImportError:
            raise ParseError("pdfminer.six 미설치 — pip install pdfminer.six")

    if suffix == ".eml":
        import email as _email
        import email.policy
        msg = _email.message_from_bytes(p.read_bytes(), policy=email.policy.default)
        parts: list[str] = []
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    payload = part.get_payload(decode=True)
                    if payload:
                        parts.append(payload.decode("utf-8", errors="replace"))
        else:
            payload = msg.get_payload(decode=True)
            if payload:
                parts.append(payload.decode("utf-8", errors="replace"))
        return "\n".join(parts)

    raise ParseError(f"지원하지 않는 파일 형식: {suffix} (지원: .txt .docx .pdf .md .eml)")


def extract_text_from_file_bytes(file_bytes: bytes, suffix: str) -> str:
    """파일 바이트 → 텍스트 추출. 백엔드 multipart 업로드 경로 전용."""
    suffix = suffix.lower()

    if suffix in (".txt", ".md"):
        for enc in ("utf-8-sig", "utf-8", "cp949", "euc-kr"):
            try:
                return file_bytes.decode(enc)
            except (UnicodeDecodeError, LookupError):
                continue
        return file_bytes.decode("utf-8", errors="replace")

    if suffix == ".docx":
        try:
            import io
            import docx  # python-docx
            doc = docx.Document(io.BytesIO(file_bytes))
            return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        except ImportError:
            raise ParseError("python-docx 미설치 — pip install python-docx")

    if suffix == ".pdf":
        try:
            import io
            from pdfminer.high_level import extract_text as _pdf_extract
            return _pdf_extract(io.BytesIO(file_bytes)) or ""
        except ImportError:
            raise ParseError("pdfminer.six 미설치 — pip install pdfminer.six")

    if suffix == ".eml":
        import email as _email
        import email.policy
        msg = _email.message_from_bytes(file_bytes, policy=email.policy.default)
        parts: list[str] = []
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    payload = part.get_payload(decode=True)
                    if payload:
                        parts.append(payload.decode("utf-8", errors="replace"))
        else:
            payload = msg.get_payload(decode=True)
            if payload:
                parts.append(payload.decode("utf-8", errors="replace"))
        return "\n".join(parts)

    raise ParseError(f"지원하지 않는 파일 형식: {suffix} (지원: .txt .docx .pdf .md .eml)")


# ── LLM 프롬프트 ─────────────────────────────────────────────────────────────

_PROMPT_TEMPLATE = """당신은 한국 중소기업 직원의 비서입니다. 받은 메일/메시지에서 사용자가 무엇을 해야 하는지 명확히 정리해야 합니다.

[받은 메시지]
{text}

[추출 규칙]
1. 한국 비즈니스 관용 표현을 정확히 이해하세요:
   - "협조 부탁드립니다" = 정중한 요청 (거절하기 어려움)
   - "검토해주시면 감사하겠습니다" = 검토 + 회신 요청
   - "가능하시다면" = 선택적 요청 (P2 또는 P3)

2. 액션은 명확한 동사로 표현하세요.

3. 우선순위: P1(긴급/필수) P2(권장) P3(선택)

4. 마감일이 없으면 parsed_date를 null로.

5. 발신자 의도는 1~2 문장으로.

[반드시 다음 JSON 형식으로만 응답하세요. 다른 텍스트 금지]
{{
  "actions": [{{"text": "액션 내용", "priority": "P1|P2|P3", "rationale": "이유"}}],
  "deadline": {{"raw_text": "원문 마감 표현", "parsed_date": "YYYY-MM-DD 또는 null", "confidence": 0.0~1.0, "time_text": "시간 표현 또는 빈 문자열"}},
  "required_materials": [{{"name": "자료명", "is_optional": false, "rationale": "이유"}}],
  "intent": {{"summary": "의도 요약", "tone": "formal|informal|urgent|casual", "expected_response": "기대 응답"}},
  "confidence": 0.0~1.0
}}"""


def _build_prompt(text: str) -> str:
    return _PROMPT_TEMPLATE.format(text=text[:8000])


# ── 규칙 기반 휴리스틱 파서 (LLM 미로드 시 fallback) ──────────────────────────

def _heuristic_parse(text: str, today: Optional[date] = None) -> ParsedResult:
    """간단한 규칙 기반 파싱 — LLM 미로드 환경(테스트 포함)에서 사용."""
    result = ParsedResult()
    today = today or date.today()

    # 마감일 추출
    date_str = parse_korean_date(text, today)
    time_str = parse_korean_time(text)
    raw_date_text = ""
    for pattern in [r"다음\s*주\s*\S+요일?", r"이번\s*주\s*\S+요일?", r"\d+월\s*\d+일",
                    r"\d+일까지", r"내일", r"모레", r"월말", r"이번\s*달\s*말"]:
        m = re.search(pattern, text)
        if m:
            raw_date_text = m.group(0)
            break
    result.deadline = Deadline(
        raw_text=raw_date_text,
        parsed_date=date_str,
        confidence=0.85 if date_str else 0.3,
        time_text=time_str,
    )

    # 액션 추출 — 동사 패턴
    action_patterns = [
        (r"([\w\s]+(?:작성|제출|준비|검토|확인|발송|전달|정리|취합|보고|답변|회신)\s*(?:해|해주|드리|부탁|요청)?)", "P1"),
        (r"([\w\s]+(?:첨부|첨부해|첨부해주)\s*(?:주시면|주세요)?)", "P2"),
    ]
    seen: set[str] = set()
    for pattern, priority in action_patterns:
        for m in re.finditer(pattern, text):
            action_text = m.group(1).strip()[:200]
            if action_text and action_text not in seen and len(action_text) > 5:
                seen.add(action_text)
                result.actions.append(ActionItem(text=action_text, priority=priority))

    if not result.actions:
        # 문장 단위 분리 — 요청성 문장 추출
        sentences = re.split(r"[.!?。\n]+", text)
        for sent in sentences:
            sent = sent.strip()
            if len(sent) > 10 and re.search(r"부탁|요청|주세요|주시면|드립니다|드리겠|바랍니다", sent):
                result.actions.append(ActionItem(
                    text=sent[:200],
                    priority="P1",
                    rationale="요청 패턴 감지",
                ))
                if len(result.actions) >= 3:
                    break

    # 필요 자료 추출 — 명사형 자료명만 추출 (동사형 "첨부"/"별첨" 접미사 제외)
    # 자료 뒤에 오는 문맥이 액션 요청(검토/확인 등)이면 제외, 전달 요청(첨부/보내 등)이면 포함
    _mat_pattern = re.compile(
        r"([\w\s]{1,20}(?:자료|데이터|파일|문서|견적|보고서|서류|서식|양식|명세서|청구서|영수증|목록|리스트|표|내역|기안|계획서|제안서))"
    )
    _action_ctx_re = re.compile(
        r"(?:작성|제출|확인|검토|날인|서명|완료|처리|수정|진행)\s*(?:해|해주|드리|부탁|요청|바랍|주세요|주시면|합니다|드립니다)"
    )
    _delivery_ctx_re = re.compile(r"첨부|보내|전달|제공|공유")
    seen_mat: set[str] = set()
    for m in _mat_pattern.finditer(text):
        name = m.group(1).strip()[:40]
        if not name or len(name) <= 3:
            continue
        ctx = text[m.end():min(len(text), m.end() + 50)]
        if _action_ctx_re.search(ctx) and not _delivery_ctx_re.search(ctx):
            continue
        if name in seen_mat:
            continue
        seen_mat.add(name)
        is_optional = bool(re.search(r"가능하시다면|여유가 되시면|선택", text[m.start()-20:m.start()+20]))
        result.required_materials.append(MaterialItem(name=name, is_optional=is_optional))
        if len(result.required_materials) >= 5:
            break

    # 의도 + 톤
    is_urgent = bool(re.search(r"긴급|즉시|바로|오늘|내일|급", text))
    is_formal = bool(re.search(r"드립니다|드리겠습니다|주십시오|협조|감사합니다", text))
    tone = "urgent" if is_urgent else ("formal" if is_formal else "informal")

    _greeting = re.compile(r"^(안녕|수고|감사|죄송|좋은|바쁜|처음)")
    _request_kw = re.compile(r"부탁|요청|드립니다|바랍니다|주세요|해주시면|검토|제출|확인|보내|작성|첨부")
    sentences = [s.strip() for s in re.split(r"[.!?。\n]+", text) if len(s.strip()) > 10]
    non_greet = [s for s in sentences if not _greeting.search(s)]
    request_sents = [s for s in non_greet if _request_kw.search(s)]
    _summary_src = request_sents[0] if request_sents else (non_greet[0] if non_greet else (sentences[0] if sentences else ""))
    summary = _summary_src[:60] if _summary_src else "요청 메시지 분석 완료"
    result.intent = Intent(summary=summary, tone=tone, expected_response="회신 또는 자료 제출")

    # 신뢰도
    result.confidence = 0.82 if result.actions else 0.55

    return result


def _parse_llm_response(raw: str, text: str, today: Optional[date] = None) -> ParsedResult:
    """LLM 응답 JSON 파싱 + 스키마 검증 + fallback."""
    raw = raw.strip()
    # JSON 블록 추출
    m = re.search(r"\{.*\}", raw, re.DOTALL)
    if not m:
        return _heuristic_parse(text, today)

    try:
        data = json.loads(m.group(0))
    except json.JSONDecodeError:
        return _heuristic_parse(text, today)

    errors = validate_schema(data)
    if errors:
        return _heuristic_parse(text, today)

    result = ParsedResult.from_dict(data)

    # 날짜 보강 — LLM이 한국어 날짜를 ISO로 변환 못 했을 때
    if result.deadline.raw_text and not result.deadline.parsed_date:
        result.deadline.parsed_date = parse_korean_date(result.deadline.raw_text, today)
    # 시간 보강 — LLM이 time_text를 채우지 않은 경우 (raw_text 대신 원문 전체에서 탐색)
    if not result.deadline.time_text:
        result.deadline.time_text = parse_korean_time(text)

    return result


# ── 신뢰도 분기 §6.1 ─────────────────────────────────────────────────────────

def confidence_level(confidence: float) -> str:
    """신뢰도 수치 → 레벨 문자열."""
    if confidence >= 0.90:
        return "high"
    if confidence >= 0.70:
        return "medium"
    if confidence >= 0.50:
        return "low"
    return "failed"


# ── 메인 진입점 ───────────────────────────────────────────────────────────────

def parse_text(
    text: str,
    input_format: str = "text",
    llm: Any = None,
    today: Optional[date] = None,
) -> ParsedResult:
    """
    텍스트 → ParsedResult.

    Args:
        text: 입력 텍스트
        input_format: text | email | docx | pdf | md
        llm: LlmRuntime 인스턴스 (None이면 휴리스틱 사용)
        today: 날짜 기준 (테스트 주입용)
    """
    # 길이 제약
    if len(text) < MIN_TEXT_LENGTH:
        raise TextTooShortError(
            f"메시지가 너무 짧습니다 (최소 {MIN_TEXT_LENGTH}자, 현재 {len(text)}자)"
        )
    if len(text) > MAX_TEXT_LENGTH:
        raise TextTooLongError(
            f"메시지가 너무 깁니다 (최대 {MAX_TEXT_LENGTH:,}자, 현재 {len(text):,}자)"
        )

    # PII 마스킹
    masked = mask_pii(text)

    # LLM 또는 휴리스틱
    result: ParsedResult
    if llm is not None and getattr(llm, "status", "") == "ready":
        try:
            prompt = _build_prompt(masked)
            raw_output = ""
            for token in llm.generate(prompt, max_tokens=1024, stop=["<|im_end|>", "</s>"]):
                raw_output += token
            result = _parse_llm_response(raw_output, masked, today)
        except Exception:
            result = _heuristic_parse(masked, today)
    else:
        result = _heuristic_parse(masked, today)

    result.masked_text = masked
    result.input_format = input_format
    return result
