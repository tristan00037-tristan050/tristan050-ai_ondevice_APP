from __future__ import annotations

import csv
import hashlib
import json
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, Iterator, Optional


@dataclass
class RawRecord:
    source_path: str
    content: str
    file_type: str
    sha256: str
    metadata: dict = field(default_factory=dict)


@dataclass
class CollectStats:
    total_files: int = 0
    by_type: Dict[str, int] = field(default_factory=dict)
    duplicates_skipped: int = 0
    failed_files: int = 0
    failed_paths: list[str] = field(default_factory=list)


class DataCollector:
    """Collect supported source documents into normalized raw text records."""

    SUPPORTED = {".pdf", ".docx", ".txt", ".jsonl", ".json", ".csv"}

    def __init__(self, dedup_cache: str = "tmp/pipeline_dedup.json"):
        self.dedup_path = Path(dedup_cache)
        self._seen: set[str] = self._load()
        self.stats = CollectStats()

    def _load(self) -> set[str]:
        if not self.dedup_path.exists():
            return set()
        try:
            data = json.loads(self.dedup_path.read_text(encoding="utf-8"))
            if isinstance(data, list):
                return {str(x) for x in data}
        except Exception:
            pass
        return set()

    def _save(self) -> None:
        self.dedup_path.parent.mkdir(parents=True, exist_ok=True)
        ordered = sorted(self._seen)
        self.dedup_path.write_text(
            json.dumps(ordered, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    @staticmethod
    def _sha256(text: str) -> str:
        return hashlib.sha256(text.encode("utf-8")).hexdigest()

    @staticmethod
    def _read_text_file(path: Path) -> str:
        return path.read_text(encoding="utf-8", errors="ignore")

    def _parse_pdf(self, path: Path) -> Optional[str]:
        try:
            import pdfplumber
        except Exception:
            return None
        try:
            with pdfplumber.open(str(path)) as pdf:
                pages = [(page.extract_text() or "") for page in pdf.pages]
            return "\n".join(pages).strip()
        except Exception:
            return None

    def _parse_docx(self, path: Path) -> Optional[str]:
        try:
            from docx import Document
        except Exception:
            return None
        try:
            doc = Document(str(path))
            parts: list[str] = []
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    parts.append(paragraph.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if values:
                        parts.append(" | ".join(values))
            return "\n".join(parts).strip()
        except Exception:
            return None

    def _parse_jsonl(self, path: Path) -> Optional[str]:
        try:
            lines: list[str] = []
            for raw_line in path.read_text(encoding="utf-8", errors="ignore").splitlines():
                stripped = raw_line.strip()
                if not stripped:
                    continue
                try:
                    obj = json.loads(stripped)
                except Exception:
                    lines.append(stripped)
                else:
                    lines.append(json.dumps(obj, ensure_ascii=False, sort_keys=True))
            return "\n".join(lines).strip()
        except Exception:
            return None

    def _parse_json(self, path: Path) -> Optional[str]:
        try:
            obj = json.loads(path.read_text(encoding="utf-8", errors="ignore"))
            return json.dumps(obj, ensure_ascii=False, sort_keys=True, indent=2).strip()
        except Exception:
            return None

    def _parse_csv(self, path: Path) -> Optional[str]:
        try:
            rows: list[str] = []
            with path.open("r", encoding="utf-8", errors="ignore", newline="") as fh:
                reader = csv.reader(fh)
                for row in reader:
                    rows.append(",".join(cell.strip() for cell in row))
            return "\n".join(rows).strip()
        except Exception:
            return None

    def _parse(self, path: Path) -> Optional[str]:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._parse_pdf(path)
        if suffix == ".docx":
            return self._parse_docx(path)
        if suffix == ".txt":
            try:
                return self._read_text_file(path).strip()
            except Exception:
                return None
        if suffix == ".jsonl":
            return self._parse_jsonl(path)
        if suffix == ".json":
            return self._parse_json(path)
        if suffix == ".csv":
            return self._parse_csv(path)
        return None

    def collect(self, source_dir: str) -> Iterator[RawRecord]:
        for fp in sorted(Path(source_dir).rglob("*")):
            if not fp.is_file():
                continue
            suffix = fp.suffix.lower()
            if suffix not in self.SUPPORTED:
                continue
            self.stats.total_files += 1
            file_type = suffix.lstrip(".")
            self.stats.by_type[file_type] = self.stats.by_type.get(file_type, 0) + 1

            content = self._parse(fp)
            if content is None:
                self.stats.failed_files += 1
                self.stats.failed_paths.append(str(fp))
                continue

            sha = self._sha256(content)
            if sha in self._seen:
                self.stats.duplicates_skipped += 1
                continue

            self._seen.add(sha)
            yield RawRecord(
                source_path=str(fp),
                content=content,
                file_type=file_type,
                sha256=sha,
                metadata={},
            )

        self._save()

    def get_stats(self) -> dict:
        return {
            "total_files": self.stats.total_files,
            "by_type": dict(self.stats.by_type),
            "duplicates_skipped": self.stats.duplicates_skipped,
            "failed_files": self.stats.failed_files,
            "failed_paths": list(self.stats.failed_paths),
        }
