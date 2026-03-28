from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from scripts.pipeline.collector_v2 import DataCollector


def test_collect_supported_files_and_dedup(tmp_path):
    source = tmp_path / "src"
    source.mkdir()

    (source / "a.txt").write_text("한국어 테스트 문장입니다.\n둘째 줄입니다.", encoding="utf-8")
    (source / "dup.txt").write_text("한국어 테스트 문장입니다.\n둘째 줄입니다.", encoding="utf-8")
    (source / "data.json").write_text(json.dumps({"a": 1, "b": "값"}, ensure_ascii=False), encoding="utf-8")
    (source / "data.csv").write_text("col1,col2\n값1,값2\n", encoding="utf-8")

    doc = Document()
    doc.add_paragraph("문서 첫 줄")
    doc.add_paragraph("문서 둘째 줄")
    doc.add_paragraph("문서 셋째 줄")
    doc.save(str(source / "sample.docx"))

    dedup_cache = tmp_path / "dedup.json"
    collector = DataCollector(dedup_cache=str(dedup_cache))
    records = list(collector.collect(str(source)))
    stats = collector.get_stats()

    assert stats["total_files"] == 5
    assert stats["duplicates_skipped"] == 1
    assert len(records) == 4
    assert dedup_cache.exists()
    assert any(record.file_type == "docx" for record in records)
    assert any(record.file_type == "json" for record in records)


def test_failed_file_is_tracked(tmp_path, monkeypatch):
    source = tmp_path / "src"
    source.mkdir()
    bad = source / "bad.txt"
    bad.write_text("실패 유도", encoding="utf-8")

    collector = DataCollector(dedup_cache=str(tmp_path / "dedup.json"))

    def broken_parse(path):
        return None

    monkeypatch.setattr(collector, "_parse", broken_parse)
    records = list(collector.collect(str(source)))
    stats = collector.get_stats()

    assert records == []
    assert stats["failed_files"] == 1
    assert str(bad) in stats["failed_paths"]
